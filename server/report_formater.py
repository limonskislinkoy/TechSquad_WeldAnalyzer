from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import os
from collections import defaultdict
import math

CLASS_LABELS = {
    1: ("пора", "A"),
    2: ("включение", "B"),
    3: ("подрез", "Fc"),
    4: ("прожог", "Fd"),
    5: ("трещина", "E"),
    6: ("наплыв", "G"),
    7: ("эталон1", "n/a"),
    8: ("эталон2", "n/a"),
    9: ("эталон3", "n/a"),
    10: ("пора-скрытая", "A"),
    11: ("утяжина", "Fa"),
    12: ("несплавление", "Db"),
    13: ("непровар корня", "Da"),
}

IMAGE_WIDTH_PX = 31500
IMAGE_HEIGHT_PX = 1152
IMAGE_WIDTH_MM = 3100.0
IMAGE_HEIGHT_MM = (IMAGE_HEIGHT_PX / IMAGE_WIDTH_PX) * IMAGE_WIDTH_MM

include_reason = True  # Включить добавление причины в заключение


def get_defect_description(defect_class):
    return CLASS_LABELS.get(defect_class, ("неизвестно", "n/a"))


def calculate_defect_size(points):
    if not points or len(points) < 4:
        return 0, 0
    x_coords = points[::2]
    y_coords = points[1::2]
    min_x, max_x = min(x_coords), max(x_coords)
    min_y, max_y = min(y_coords), max(y_coords)
    width = (max_x - min_x) * IMAGE_WIDTH_MM
    height = (max_y - min_y) * IMAGE_HEIGHT_MM
    return width, height


def calculate_defect_area(points):
    width, height = calculate_defect_size(points)
    return math.pi * width * height / 4


def format_defect_description(code, width, height, count, is_pore=False):
    if is_pore:
        diameter = max(width, height)
        return (
            f"{code}{int(diameter)}<"
            if count == 1
            else f"{count}{code}{int(diameter)}<"
        )
    return (
        f"{code}{int(width)}x{int(height)}<"
        if count == 1
        else f"{count}{code}{int(width)}x{int(height)}<"
    )


def get_conclusion(defects_data):
    if not defects_data:
        return "годен", ""

    total_pore_area = 0
    total_root_length = 0

    for cls, value in defects_data:
        if cls in [1, 10] and isinstance(value, float):
            total_pore_area += value
        elif cls == 13 and isinstance(value, tuple):
            width, _ = value
            total_root_length += width

    for cls, value in defects_data:
        if not isinstance(value, tuple):
            continue
        width, _ = value
        if cls == 2 and width > 5:
            return "ремонт", "включение > 5 мм"
        if cls == 11 and width > 1:
            return "ремонт", "утяжина > 1 мм"
        if cls == 3 and width > 1:
            return "ремонт", "подрез > 1 мм"

    if total_pore_area > 30:
        return "ремонт", f"суммарная площадь пор {int(total_pore_area)} мм²"
    if total_root_length > 30:
        return "ремонт", f"непровар корня > 30 мм"

    return "годен", ""


def create_report(image_name, user_id, polygons, intervals, output_folder):
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    def add_centered_paragraph(text, bold=False, underline=False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.underline = underline
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    add_centered_paragraph(f"ЗАКЛЮЧЕНИЕ № {image_name}", bold=True)
    add_centered_paragraph(f"от {datetime.now().strftime('%d.%m.%Y')} года", bold=True)
    add_centered_paragraph(
        "по результатам контроля качества сварных соединений радиографическим методом",
        bold=True,
    )
    p = doc.add_paragraph()
    run = p.add_run("Тип источника ионизирующего излучения: ")
    run.bold = True
    run2 = p.add_run("рентгеновский дефектоскоп непрерывного действия")
    run2.bold = True
    run2.underline = True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    add_centered_paragraph(
        "Номер операционной технологической карты контроля: ТК-РК-ЦР-С-1020х17П-00890001",
        bold=True,
    )
    doc.add_paragraph("")

    p = doc.add_paragraph("РЕЗУЛЬТАТЫ КОНТРОЛЯ")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    table = doc.add_table(rows=1, cols=8)
    table.style = "Table Grid"

    headers = [
        "Номер сварного соединения по журналу сварки",
        "Диаметр и толщина стенки трубы, мм",
        "Шифр бригады или клеймо сварщика",
        "Номер участка контроля (координаты мерного пояса)",
        "Чувствительность контроля, мм",
        "Описание выявленных дефектов",
        "Координаты недопустимых дефектов",
        "Заключение (годен, ремонт, вырезать)",
    ]

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].alignment = 1
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    row_count = 0

    for interval, indices in intervals.items():
        row_cells = table.add_row().cells

        if row_count == 0:
            row_cells[0].text = image_name
            row_cells[1].text = "1020×17"
            row_cells[2].text = user_id if user_id != "0" else "XXX"
        else:
            row_cells[0].text = ""
            row_cells[1].text = ""
            row_cells[2].text = ""

        row_cells[3].text = interval
        row_cells[4].text = "0,50"

        defects_data = []
        grouped_defects = defaultdict(int)

        descriptions = []

        for idx in indices:
            if idx >= len(polygons):
                continue
            item = polygons[idx]
            defect_class = item.get("class")
            points = item.get("polygon", [])

            desc, code = get_defect_description(defect_class)
            if code == "n/a":
                continue

            width, height = calculate_defect_size(points)
            width = int(round(width))  # округление до целого и преобразование в int
            height = int(round(height))

            if defect_class in [1, 10]:
                area = calculate_defect_area(points)
                defects_data.append((defect_class, area))
            else:
                defects_data.append((defect_class, (width, height)))

            grouped_defects[(code, width, height)] += 1

        for (code, width, height), count in grouped_defects.items():
            is_pore = code == "A"
            desc = format_defect_description(code, width, height, count, is_pore)
            descriptions.append(desc)

        row_cells[5].text = "; ".join(descriptions) if descriptions else "н/д"

        conclusion, reason = get_conclusion(defects_data)
        row_cells[6].text = interval if conclusion == "ремонт" else "н/п"
        row_cells[7].text = (
            f"{conclusion} ({reason})" if include_reason and reason else conclusion
        )

        for cell in row_cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = 1
                for run in paragraph.runs:
                    run.font.size = Pt(10)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        row_count += 1

    if row_count > 1:
        for col in range(3):
            start_cell = table.cell(1, col)
            for row in range(2, row_count + 1):
                start_cell.merge(table.cell(row, col))

    os.makedirs(output_folder, exist_ok=True)
    output_path = os.path.join(output_folder, f"{image_name}_report.docx")
    doc.save(output_path)

    return output_path
